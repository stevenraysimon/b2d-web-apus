from flask import Flask, request, render_template, send_file, url_for
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches
from io import BytesIO
import os

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'static/uploads'

if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

all_extracted_content = []
banner_images = []

def extract_content_from_html(html_content, is_overview):
    soup = BeautifulSoup(html_content, 'html.parser')

    if is_overview:
        iframe = soup.find('iframe', {'srcdoc': True})
        if iframe:
            srcdoc_content = iframe['srcdoc']
            soup = BeautifulSoup(srcdoc_content, 'html.parser')

    # Extract content based on tag names
    extracted_content = []
    for tag in soup.find_all(['h1', 'h2', 'p', 'ul', 'li']):
        if tag.name == 'li':
            extracted_content.append(f"• {tag.get_text()}")
        elif tag.name == 'ul':
            extracted_content.append("")
        else:
            extracted_content.append(tag.get_text())

    return '\n'.join(extracted_content)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    file = request.files['file']
    is_overview = request.form.get('isOverview') == 'true'
    image_file = request.files.get('image')

    if file:
        html_content = file.read().decode('utf-8')
        text_content = extract_content_from_html(html_content, is_overview)
        all_extracted_content.append(text_content)

        image_url = None
        if image_file:
            image_path = os.path.join(app.config['UPLOAD_FOLDER'], image_file.filename)
            image_file.save(image_path)
            banner_images.append(image_path)
            image_url = url_for('static', filename=f'uploads/{image_file.filename}')
        else:
            banner_images.append(None)

        return {'content': text_content, 'image_url': image_url}, 200
    return {'error': 'File upload failed'}, 400

@app.route('/download', methods=['GET'])
def download_file():
    if all_extracted_content:
        document = Document()
        for index, content in enumerate(all_extracted_content):
            if index > 0 and banner_images[index - 1]:
                document.add_picture(banner_images[index - 1], width=Inches(6))
            document.add_paragraph(content)

        file_stream = BytesIO()
        document.save(file_stream)
        file_stream.seek(0)

        return send_file(file_stream, as_attachment=True, download_name='output.docx')
    return {'error': 'No content to download'}, 400

def clear_uploads_folder():
    folder = app.config['UPLOAD_FOLDER']
    for filename in os.listdir(folder):
        file_path = os.path.join(folder, filename)
        try:
            if os.path.isfile(file_path):
                os.unlink(file_path)
        except Exception as e:
            print(f"Error deleting file: {e}")

@app.route('/clear-content', methods=['POST'])
def clear_content():
    global all_extracted_content, banner_images
    all_extracted_content = []
    banner_images = []
    clear_uploads_folder()
    return 'Content cleared', 200

if __name__ == '__main__':
    app.run(debug=True)
